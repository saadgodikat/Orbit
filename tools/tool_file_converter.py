import os
import shutil
from typing import Dict, Any
from tools.base import BaseTool
from core.state_manager import StateManager


class FileConverterTool(BaseTool):
    """Universal file format converter supporting images, documents, and data files."""

    # Supported conversion map: source_ext → [target_exts]
    SUPPORTED = {
        # Image conversions (via Pillow)
        ".jpg":  [".png", ".bmp", ".webp", ".tiff", ".gif", ".ico"],
        ".jpeg": [".png", ".bmp", ".webp", ".tiff", ".gif", ".ico"],
        ".png":  [".jpg", ".jpeg", ".bmp", ".webp", ".tiff", ".gif", ".ico"],
        ".bmp":  [".png", ".jpg", ".webp", ".tiff", ".gif"],
        ".webp": [".png", ".jpg", ".bmp", ".tiff", ".gif"],
        ".tiff": [".png", ".jpg", ".bmp", ".webp", ".gif"],
        ".gif":  [".png", ".jpg", ".bmp", ".webp"],
        ".ico":  [".png", ".jpg"],
        # Document conversions
        ".pdf":  [".docx"],
        ".docx": [".pdf", ".txt"],
        ".md":   [".html", ".txt"],
        ".html": [".pdf", ".txt"],
        # Data conversions
        ".csv":  [".json", ".txt"],
        ".json": [".csv", ".txt"],
        ".txt":  [".md"],
    }

    @classmethod
    def get_name(cls) -> str:
        return "file_converter"

    @classmethod
    def get_description(cls) -> str:
        return (
            "TOOL NAME: file_converter\n"
            "DESCRIPTION: Converts files between formats. Supports:\n"
            "  - Images: JPG, PNG, BMP, WEBP, TIFF, GIF, ICO (any to any)\n"
            "  - Documents: PDF→DOCX, DOCX→PDF/TXT, MD→HTML/TXT, HTML→PDF\n"
            "  - Data: CSV↔JSON, TXT↔MD\n"
            "REQUIRED ARGS:\n"
            "  - input_path (str): Absolute path to the source file\n"
            "  - output_format (str): Target format extension, e.g. '.png', '.docx', '.json'\n"
            "OPTIONAL ARGS:\n"
            "  - output_path (str): Custom output path (default: same dir, new extension)\n"
            "STATE OUTPUT: Saves converted file path to {{converter_last_output}}"
        )

    def _convert_image(self, input_path: str, output_path: str, target_ext: str) -> bool:
        """Convert between image formats using Pillow."""
        from PIL import Image

        img = Image.open(input_path)

        # Handle transparency: convert RGBA to RGB for formats that don't support alpha
        if target_ext in [".jpg", ".jpeg", ".bmp", ".ico"] and img.mode in ("RGBA", "P"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            bg.paste(img, mask=img.split()[3])
            img = bg

        # Handle ICO: resize to standard icon size
        if target_ext == ".ico":
            img = img.resize((256, 256), Image.LANCZOS)

        # Save with format-specific options
        save_kwargs = {}
        if target_ext in [".jpg", ".jpeg"]:
            save_kwargs["quality"] = 95
            save_kwargs["optimize"] = True
            img = img.convert("RGB")
        elif target_ext == ".png":
            save_kwargs["optimize"] = True
        elif target_ext == ".webp":
            save_kwargs["quality"] = 90

        img.save(output_path, **save_kwargs)
        return True

    def _convert_pdf_to_docx(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to DOCX using pdf2docx."""
        try:
            from pdf2docx import Converter
        except ImportError:
            print("\033[38;5;214m[ INSTALL ]\033[0m Installing pdf2docx...")
            import subprocess
            subprocess.run(["pip", "install", "pdf2docx"], capture_output=True)
            from pdf2docx import Converter

        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True

    def _convert_docx_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX to PDF via HTML intermediate using Playwright."""
        try:
            from docx import Document
        except ImportError:
            print("\033[38;5;214m[ INSTALL ]\033[0m Installing python-docx...")
            import subprocess
            subprocess.run(["pip", "install", "python-docx"], capture_output=True)
            from docx import Document

        # Read DOCX content
        doc = Document(input_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name.lower()
                if "heading 1" in style:
                    paragraphs.append(f"<h1>{text}</h1>")
                elif "heading 2" in style:
                    paragraphs.append(f"<h2>{text}</h2>")
                elif "heading 3" in style:
                    paragraphs.append(f"<h3>{text}</h3>")
                else:
                    paragraphs.append(f"<p>{text}</p>")

        html_content = f"""<!DOCTYPE html>
<html><head><style>
body {{ font-family: Georgia, serif; padding: 40px; line-height: 1.8; color: #1a1a1a; max-width: 800px; margin: 0 auto; }}
h1, h2, h3 {{ margin-top: 1.5em; }}
</style></head><body>{"".join(paragraphs)}</body></html>"""

        # Write temp HTML, then convert to PDF via Playwright
        tmp_html = output_path.replace(".pdf", "_tmp.html")
        with open(tmp_html, "w", encoding="utf-8") as f:
            f.write(html_content)

        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(f"file://{os.path.abspath(tmp_html)}", wait_until="networkidle")
            page.pdf(
                path=output_path,
                format="A4",
                margin={"top": "20mm", "right": "15mm", "bottom": "20mm", "left": "15mm"},
                print_background=True,
            )
            browser.close()

        os.remove(tmp_html)
        return True

    def _convert_docx_to_txt(self, input_path: str, output_path: str) -> bool:
        """Extract plain text from DOCX."""
        try:
            from docx import Document
        except ImportError:
            import subprocess
            subprocess.run(["pip", "install", "python-docx"], capture_output=True)
            from docx import Document

        doc = Document(input_path)
        text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        return True

    def _convert_md_to_html(self, input_path: str, output_path: str) -> bool:
        """Convert Markdown to HTML."""
        import markdown

        with open(input_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        html = markdown.markdown(md_content, extensions=["tables", "fenced_code", "nl2br", "sane_lists"])
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><style>
body {{ font-family: Georgia, serif; padding: 40px; max-width: 800px; margin: 0 auto; line-height: 1.8; color: #1a1a1a; }}
table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background: #f5f5f5; }}
code {{ background: #f3f3f3; padding: 2px 5px; border-radius: 3px; }}
pre {{ background: #f7f7f7; padding: 16px; border-radius: 3px; overflow-x: auto; }}
</style></head><body>{html}</body></html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        return True

    def _convert_html_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert HTML to PDF via Playwright."""
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(f"file://{os.path.abspath(input_path)}", wait_until="networkidle")
            page.pdf(
                path=output_path,
                format="A4",
                margin={"top": "20mm", "right": "15mm", "bottom": "20mm", "left": "15mm"},
                print_background=True,
            )
            browser.close()
        return True

    def _convert_csv_to_json(self, input_path: str, output_path: str) -> bool:
        """Convert CSV to JSON."""
        import csv
        import json

        with open(input_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(rows, f, indent=2, ensure_ascii=False)
        return True

    def _convert_json_to_csv(self, input_path: str, output_path: str) -> bool:
        """Convert JSON array to CSV."""
        import csv
        import json

        with open(input_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        if not isinstance(data, list) or not data:
            print("\033[38;5;196m[ ERROR ]\033[0m JSON must be an array of objects for CSV conversion.")
            return False

        # Flatten nested dicts to strings
        flat_data = []
        for row in data:
            flat_row = {}
            for k, v in row.items():
                flat_row[k] = str(v) if isinstance(v, (dict, list)) else v
            flat_data.append(flat_row)

        fieldnames = list(flat_data[0].keys())
        with open(output_path, "w", encoding="utf-8", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(flat_data)
        return True

    def _convert_to_txt(self, input_path: str, output_path: str) -> bool:
        """Extract text content from any readable file."""
        with open(input_path, "r", encoding="utf-8") as f:
            content = f.read()
        # Strip HTML tags if present
        import re
        if "<html" in content.lower() or "<body" in content.lower():
            content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL)
            content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL)
            content = re.sub(r"<[^>]+>", "", content)
            content = re.sub(r"\n{3,}", "\n\n", content)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content.strip())
        return True

    def run(self, state: StateManager, args: Dict[str, Any]) -> bool:
        input_path = args.get("input_path")
        output_format = args.get("output_format", "").lower()
        output_path = args.get("output_path", "")

        if not input_path:
            print("\033[38;5;196m[ ERROR ]\033[0m file_converter requires 'input_path'.")
            return False

        if not os.path.exists(input_path):
            print(f"\033[38;5;196m[ ERROR ]\033[0m File not found: {input_path}")
            return False

        if not output_format.startswith("."):
            output_format = "." + output_format

        source_ext = os.path.splitext(input_path)[1].lower()

        # Validate conversion is supported
        if source_ext not in self.SUPPORTED:
            print(f"\033[38;5;196m[ ERROR ]\033[0m Unsupported source format: {source_ext}")
            return False

        if output_format not in self.SUPPORTED[source_ext]:
            supported = ", ".join(self.SUPPORTED[source_ext])
            print(f"\033[38;5;196m[ ERROR ]\033[0m Cannot convert {source_ext} → {output_format}. Supported: {supported}")
            return False

        # Generate output path if not specified
        if not output_path:
            base = os.path.splitext(input_path)[0]
            output_path = base + output_format

        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        print(f"\033[38;5;39m[ CONVERTER ]\033[0m {source_ext} → {output_format}")
        print(f"\033[38;5;244m  Input  : {input_path}\033[0m")
        print(f"\033[38;5;244m  Output : {output_path}\033[0m")

        try:
            # Image conversions
            img_exts = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tiff", ".gif", ".ico"}
            if source_ext in img_exts and output_format in img_exts:
                self._convert_image(input_path, output_path, output_format)

            # PDF → DOCX
            elif source_ext == ".pdf" and output_format == ".docx":
                self._convert_pdf_to_docx(input_path, output_path)

            # DOCX → PDF
            elif source_ext == ".docx" and output_format == ".pdf":
                self._convert_docx_to_pdf(input_path, output_path)

            # DOCX → TXT
            elif source_ext == ".docx" and output_format == ".txt":
                self._convert_docx_to_txt(input_path, output_path)

            # Markdown → HTML
            elif source_ext == ".md" and output_format == ".html":
                self._convert_md_to_html(input_path, output_path)

            # HTML → PDF
            elif source_ext == ".html" and output_format == ".pdf":
                self._convert_html_to_pdf(input_path, output_path)

            # CSV → JSON
            elif source_ext == ".csv" and output_format == ".json":
                self._convert_csv_to_json(input_path, output_path)

            # JSON → CSV
            elif source_ext == ".json" and output_format == ".csv":
                self._convert_json_to_csv(input_path, output_path)

            # Anything → TXT
            elif output_format == ".txt":
                self._convert_to_txt(input_path, output_path)

            # TXT → MD (just copy with .md extension)
            elif source_ext == ".txt" and output_format == ".md":
                shutil.copy2(input_path, output_path)

            # MD → TXT (just copy with .txt extension)
            elif source_ext == ".md" and output_format == ".txt":
                shutil.copy2(input_path, output_path)

            else:
                print(f"\033[38;5;196m[ ERROR ]\033[0m Conversion {source_ext} → {output_format} not implemented.")
                return False

            file_size = os.path.getsize(output_path)
            size_str = f"{file_size / 1024:.1f}KB" if file_size < 1048576 else f"{file_size / 1048576:.1f}MB"
            print(f"\033[38;5;114m[ SUCCESS ]\033[0m Converted to {output_path} ({size_str})")

            state.set("converter_last_output", output_path)
            state.set("converter_last_format", output_format)
            return True

        except Exception as e:
            print(f"\033[38;5;196m[ ERROR ]\033[0m Conversion failed: {e}")
            return False
